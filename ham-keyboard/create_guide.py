"""Generate the Master Resource Guide Word document for the full Gemini Voice project."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ---------------------------------------------------------------------------
# Styles
# ---------------------------------------------------------------------------
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def numbered(text):
    doc.add_paragraph(text, style="List Number")


# ===========================================================================
# TITLE PAGE
# ===========================================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Gemini Voice\nMaster Resource Guide")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(
    "Building Real-Time Voice Agents with Google Gemini 3.1 Flash Live\n"
    "Two Complete Demos: Aria Executive Assistant + E-Commerce Voice Widget"
)
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(
    "Built with Claude Code  \u2022  Gemini 3.1 Flash Live  \u2022  FastAPI  \u2022  Web Audio API\n"
    "google-genai SDK  \u2022  ClickUp API  \u2022  Google Calendar API"
)
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ===========================================================================
# TABLE OF CONTENTS
# ===========================================================================
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Project Overview",
    "2. Gemini 3.1 Flash Live \u2014 The Model",
    "3. Platform Comparison \u2014 Gemini vs OpenAI vs ElevenLabs",
    "4. Project Structure \u2014 Full File Map",
    "5. Demo 1: Aria \u2014 Executive Assistant (aria-assistant/)",
    "6. Demo 2: Apex Keyboards \u2014 Sales Widget (keyboard-demo/)",
    "7. Architecture Deep Dive \u2014 How It All Works",
    "8. Setup Guide \u2014 From Zero to Running",
    "9. How to Add Your Own Tools",
    "10. What to Tell Claude Code",
    "11. Deployment to Production",
    "12. Known Issues & Tips",
]
for item in toc_items:
    doc.add_paragraph(item)
doc.add_page_break()

# ===========================================================================
# 1. PROJECT OVERVIEW
# ===========================================================================
doc.add_heading("1. Project Overview", level=1)

doc.add_paragraph(
    "This project contains two complete demos showing how to build real-time voice agents "
    "with Google\u2019s Gemini 3.1 Flash Live API. Both use the same core architecture: a Python "
    "FastAPI server acts as a WebSocket relay between a browser and Gemini\u2019s Live API, streaming "
    "audio bidirectionally in real-time."
)

doc.add_heading("The Two Demos", level=2)

doc.add_heading("Demo 1: Aria \u2014 Executive Assistant", level=3)
doc.add_paragraph(
    "A voice-powered receptionist that manages your Google Calendar and ClickUp tasks through "
    "natural conversation. Say things like \u201cWhat\u2019s on my calendar today?\u201d or \u201cCreate a task "
    "called Review Q2 Report.\u201d Aria uses server-side tool calling \u2014 Gemini decides when to call "
    "a function, your server executes it, and sends the result back. This is the more advanced demo."
)

doc.add_heading("Demo 2: Apex Keyboards \u2014 Sales Widget", level=3)
doc.add_paragraph(
    "A dark-themed e-commerce landing page for a fictional keyboard brand with a floating mic "
    "button. Visitors can ask about products, shipping, and return policies. The agent\u2019s knowledge "
    "comes from its system prompt \u2014 no tool calling needed. This is the simpler demo, focused "
    "on showing how to embed a voice widget on any website."
)

doc.add_heading("Why Native Speech-to-Speech Matters", level=2)
doc.add_paragraph(
    "Most voice assistants use a 3-step pipeline: Speech-to-Text \u2192 LLM \u2192 Text-to-Speech. "
    "Each step adds latency and loses information (tone, emotion, hesitation). Gemini 3.1 Flash "
    "Live is natively multimodal \u2014 it processes audio directly and generates audio directly."
)
bullet("Sub-second response times (no pipeline overhead)")
bullet("The model hears tone, emphasis, and emotion \u2014 not just words")
bullet("Natural interruption support (barge-in) \u2014 talk over it and it stops immediately")
bullet("No separate STT or TTS service \u2014 just one API")

doc.add_heading("Architecture (Both Demos)", level=2)
add_code(
    "Browser (mic + speaker)\n"
    "    |  WebSocket (PCM audio as base64 JSON)\n"
    "    v\n"
    "FastAPI server (Python)\n"
    "    |  google-genai SDK (WebSocket)\n"
    "    v\n"
    "Gemini 3.1 Flash Live API\n"
    "    |  tool calls (Aria only)\n"
    "    v\n"
    "Python tool functions (ClickUp, Calendar)"
)

doc.add_page_break()

# ===========================================================================
# 2. GEMINI 3.1 FLASH LIVE
# ===========================================================================
doc.add_heading("2. Gemini 3.1 Flash Live \u2014 The Model", level=1)

doc.add_heading("What\u2019s New in 3.1 Flash Live", level=2)
bullet("Google\u2019s highest-quality audio model for real-time dialogue")
bullet("Better tonal understanding \u2014 more natural, expressive responses")
bullet("Can follow conversation threads for 2x longer than previous models")
bullet("Better at filtering background noise (traffic, TV, etc.)")
bullet("Faster response times than 2.5 Flash")
bullet("90+ languages supported with real-time language switching")
bullet("Full-duplex conversation \u2014 users can interrupt mid-sentence (barge-in)")
bullet("All audio output is watermarked to prevent misuse")
bullet("Model ID: gemini-3.1-flash-live-preview")

doc.add_heading("Session Limits", level=2)
add_table(
    ["Dimension", "Limit"],
    [
        ["Context window", "128k tokens"],
        ["Audio-only session", "15 minutes max"],
        ["Audio + video session", "2 minutes max"],
        ["Concurrent sessions (pay-as-you-go)", "Up to 1,000 per project"],
    ],
)

doc.add_heading("Pricing", level=2)
doc.add_paragraph(
    "Audio is converted at 25 tokens per second for both input and output."
)
add_table(
    ["Model", "Input (per 1M tokens)", "Output (per 1M tokens)"],
    [
        ["Gemini 3.1 Flash", "$0.50", "$3.00"],
        ["Gemini 3.1 Flash-Lite", "$0.25", "$1.50"],
        ["Gemini 3.1 Pro", "$2.00", "$12.00"],
    ],
)
doc.add_paragraph(
    "Rough cost: A 5-minute voice conversation \u2248 7,500 input tokens + 7,500 output "
    "tokens. At Flash pricing: ~$0.03 per conversation. A free tier is available through "
    "Google AI Studio with rate limits."
)

doc.add_heading("How to Get an API Key", level=2)
numbered("Go to https://aistudio.google.com/apikey")
numbered("Sign in with your Google account")
numbered("Click \u2018Create API Key\u2019 and select or create a Google Cloud project")
numbered("Copy the key \u2014 you\u2019ll put this in your .env file as GOOGLE_API_KEY")

doc.add_page_break()

# ===========================================================================
# 3. COMPARISON
# ===========================================================================
doc.add_heading("3. Platform Comparison", level=1)

add_table(
    ["", "Gemini 3.1 Flash Live", "OpenAI Realtime API", "ElevenLabs Conv. AI"],
    [
        ["Architecture", "Native speech-to-speech", "Native speech-to-speech", "STT \u2192 LLM \u2192 TTS pipeline"],
        ["Deploy complexity", "Need a WebSocket server ($5-7/mo)", "Need a WebSocket server", "One <script> tag (hosted)"],
        ["Audio input cost", "~$0.50/1M tokens (25 tok/sec)", "$0.06/min + $10/1M tokens", "$0.10/min (agent time)"],
        ["Audio output cost", "~$3.00/1M tokens (25 tok/sec)", "$0.24/min + $20/1M tokens", "Included"],
        ["5-min conversation cost", "~$0.03", "~$1.50+", "~$0.50"],
        ["Voice quality", "Good (native)", "Good (native)", "Excellent (best TTS)"],
        ["Latency", "Very low", "Very low", "Low (~200ms pipeline)"],
        ["Interruption", "Yes (native)", "Yes (native)", "Yes"],
        ["Tool calling", "Yes (synchronous)", "Yes (deterministic JSON)", "Limited"],
        ["Languages", "90+", "~30", "30+"],
        ["Voice cloning", "No", "No", "Yes"],
        ["Free tier", "Yes (rate limited)", "No", "No (trial only)"],
    ],
)

doc.add_heading("When to Use Each", level=2)
bullet("Gemini 3.1 Flash Live \u2014 Best value. Cheapest per conversation. Full tool calling. Best choice if you want control and don\u2019t mind a small server.")
bullet("OpenAI Realtime API \u2014 Best structured tool calling. Higher cost but great DX. Use if you\u2019re in the OpenAI ecosystem.")
bullet("ElevenLabs \u2014 Easiest to deploy (one script tag). Best voice quality and voice cloning. Use if you want zero backend work.")

doc.add_page_break()

# ===========================================================================
# 4. PROJECT STRUCTURE
# ===========================================================================
doc.add_heading("4. Project Structure \u2014 Full File Map", level=1)

add_code(
    "Gemini Voice/\n"
    "\u251c\u2500\u2500 .env                          # API keys (GOOGLE_API_KEY, CLICKUP_API_KEY)\n"
    "\u251c\u2500\u2500 CLAUDE.md                     # Instructions for Claude Code\n"
    "\u251c\u2500\u2500 README.md                     # Project overview & setup\n"
    "\u2502\n"
    "\u251c\u2500\u2500 aria-assistant/               # Demo 1: Executive Assistant\n"
    "\u2502   \u251c\u2500\u2500 server.py                 # FastAPI server + Gemini relay + tool execution\n"
    "\u2502   \u251c\u2500\u2500 index.html                # Single-file browser UI (mic, transcripts, tool indicator)\n"
    "\u2502   \u251c\u2500\u2500 requirements.txt          # Python dependencies\n"
    "\u2502   \u251c\u2500\u2500 client_secret_*.json      # Google OAuth credentials (not committed)\n"
    "\u2502   \u251c\u2500\u2500 token.json                # OAuth token (auto-generated on first run)\n"
    "\u2502   \u2514\u2500\u2500 tools/\n"
    "\u2502       \u251c\u2500\u2500 __init__.py\n"
    "\u2502       \u251c\u2500\u2500 clickup.py            # ClickUp API: workspaces, spaces, lists, tasks\n"
    "\u2502       \u2514\u2500\u2500 calendar.py           # Google Calendar: events, create, details\n"
    "\u2502\n"
    "\u2514\u2500\u2500 keyboard-demo/                # Demo 2: E-Commerce Voice Widget\n"
    "    \u251c\u2500\u2500 server.py                 # FastAPI server + Gemini relay (no tools)\n"
    "    \u251c\u2500\u2500 requirements.txt          # Python dependencies\n"
    "    \u2514\u2500\u2500 public/\n"
    "        \u251c\u2500\u2500 index.html            # Landing page HTML\n"
    "        \u251c\u2500\u2500 css/styles.css        # Dark theme, animations, voice widget styles\n"
    "        \u2514\u2500\u2500 js/\n"
    "            \u251c\u2500\u2500 products.js       # Product catalog data + card renderer\n"
    "            \u251c\u2500\u2500 audio-recorder.js # Mic capture \u2192 PCM16 base64\n"
    "            \u251c\u2500\u2500 audio-streamer.js # PCM16 base64 \u2192 speaker playback\n"
    "            \u251c\u2500\u2500 gemini-live-client.js # WebSocket client to local server\n"
    "            \u2514\u2500\u2500 voice-widget.js   # Mic button UI + state machine"
)

doc.add_heading("Shared: .env File", level=2)
doc.add_paragraph(
    "Both demos read from the same .env file in the Gemini Voice/ root. Only GOOGLE_API_KEY is "
    "required for the keyboard demo. Aria also needs CLICKUP_API_KEY."
)
add_code(
    "GOOGLE_API_KEY=your-gemini-api-key-here\n"
    "CLICKUP_API_KEY=your-clickup-api-key    # only for Aria"
)

doc.add_heading("Shared: CLAUDE.md", level=2)
doc.add_paragraph(
    "This file gives Claude Code context about the project when working on it. It describes "
    "the architecture, how to run each demo, the audio pipeline, and the tool execution pattern. "
    "If you open this project in Claude Code, it reads CLAUDE.md automatically."
)

doc.add_page_break()

# ===========================================================================
# 5. DEMO 1: ARIA
# ===========================================================================
doc.add_heading("5. Demo 1: Aria \u2014 Executive Assistant", level=1)

doc.add_paragraph(
    "Aria is a voice-powered receptionist that manages your schedule and tasks. It demonstrates "
    "the full power of Gemini\u2019s Live API including real-time tool calling."
)

doc.add_heading("What Aria Can Do", level=2)
bullet("\u201cWhat\u2019s on my calendar today?\u201d \u2014 reads your Google Calendar")
bullet("\u201cSchedule a meeting for Friday at 2pm\u201d \u2014 creates calendar events")
bullet("\u201cWhat workspaces do I have in ClickUp?\u201d \u2014 browses your ClickUp")
bullet("\u201cCreate a task called Review Q2 Report\u201d \u2014 creates ClickUp tasks")
bullet("\u201cUpdate that task to high priority\u201d \u2014 modifies existing tasks")
bullet("\u201cClose the task\u201d \u2014 marks tasks as complete")

doc.add_heading("aria-assistant/server.py", level=2)
doc.add_paragraph(
    "The core of Aria. A FastAPI app with a /ws WebSocket endpoint. On each connection it:"
)
numbered("Creates a Gemini Live session with LiveConnectConfig (model, system prompt, tools, transcription)")
numbered("Runs two concurrent async tasks: browser_to_gemini() and gemini_to_browser()")
numbered("browser_to_gemini(): receives base64 audio from the browser, decodes it, calls session.send_realtime_input(audio=Blob(...))")
numbered("gemini_to_browser(): receives Gemini responses and forwards audio, transcripts, turn_complete, and interrupted events")
numbered("When Gemini sends a tool_call: executes the Python function from the TOOLS dict, sends the result back via session.send_tool_response()")

doc.add_paragraph("Key configuration in server.py:")
add_table(
    ["Setting", "Value"],
    [
        ["Model", "gemini-3.1-flash-live-preview"],
        ["Port", "8000"],
        ["Response modalities", '["AUDIO"]'],
        ["System prompt", "~60 lines defining Aria\u2019s personality, capabilities, and guidelines"],
        ["Tools", "13 function declarations (9 ClickUp + 4 Calendar)"],
        ["Transcription", "Enabled for both input and output"],
    ],
)

doc.add_heading("aria-assistant/index.html", level=2)
doc.add_paragraph(
    "A single-file browser client (no build step). Contains all HTML, CSS, and JavaScript. "
    "Key features:"
)
bullet("Mic capture via AudioWorklet (inline data: URL \u2014 no separate worklet file)")
bullet("PCM16 encoding at 16kHz, sent as base64 JSON over WebSocket")
bullet("Audio playback queue with sequential scheduling at 24kHz")
bullet("clearPlayback() for instant barge-in (empties queue, stops playback)")
bullet("Real-time transcript display (user messages right-aligned, agent left-aligned)")
bullet("Tool call indicator with spinner (shows \u201cChecking today\u2019s calendar...\u201d etc.)")

doc.add_heading("aria-assistant/tools/clickup.py", level=2)
doc.add_paragraph(
    "ClickUp API v2 wrapper. 9 functions that all return {status: success/error, ...} dicts. "
    "Uses the CLICKUP_API_KEY from .env. Functions:"
)
add_table(
    ["Function", "What It Does"],
    [
        ["get_workspaces()", "Lists all workspaces (teams)"],
        ["get_spaces(team_id)", "Lists spaces in a workspace"],
        ["get_lists(space_id)", "Lists all task lists (including inside folders)"],
        ["get_tasks(list_id, status?)", "Gets tasks from a list, optional status filter"],
        ["get_task_details(task_id)", "Full details of one task"],
        ["create_task(list_id, name, ...)", "Creates a new task"],
        ["update_task(task_id, status?, priority?)", "Updates status or priority"],
        ["add_comment(task_id, comment_text)", "Adds a comment to a task"],
        ["close_task(task_id)", "Closes/completes a task"],
    ],
)

doc.add_heading("aria-assistant/tools/calendar.py", level=2)
doc.add_paragraph(
    "Google Calendar API wrapper using OAuth 2.0. Requires a client_secret JSON file from "
    "Google Cloud Console. On first run, opens a browser for Google login and saves token.json. "
    "After that, tokens auto-refresh. 4 functions:"
)
add_table(
    ["Function", "What It Does"],
    [
        ["get_todays_events()", "All events for today"],
        ["get_upcoming_events(days?)", "Events for the next N days (default 7)"],
        ["get_event_details(event_id)", "Full details of one event"],
        ["create_event(summary, start, end, ...)", "Creates a new calendar event"],
    ],
)

doc.add_heading("How Tool Calling Works", level=2)
doc.add_paragraph(
    "This is the key pattern. Tools are registered in two parallel structures in server.py:"
)
numbered("TOOLS dict \u2014 maps function name strings to Python callables: {'get_workspaces': clickup.get_workspaces, ...}")
numbered("TOOL_DECLARATIONS list \u2014 JSON schema descriptions sent to Gemini so it knows what tools exist and what parameters they take")
numbered("When you speak, Gemini decides if a tool is needed based on your request")
numbered("Gemini pauses speaking and sends a tool_call message with function name + args")
numbered("Server executes the function via execute_tool(name, args)")
numbered("Server sends the result back via session.send_tool_response(function_responses=[...])")
numbered("Gemini resumes speaking, incorporating the result into its response")
doc.add_paragraph(
    "Tool calls are synchronous \u2014 the model pauses while your function runs. The functions "
    "themselves use the requests library (synchronous HTTP), called from async context."
)

doc.add_page_break()

# ===========================================================================
# 6. DEMO 2: KEYBOARD
# ===========================================================================
doc.add_heading("6. Demo 2: Apex Keyboards \u2014 Sales Widget", level=1)

doc.add_paragraph(
    "A simpler demo focused on embedding a voice widget on a website. No tool calling \u2014 the "
    "agent\u2019s knowledge comes entirely from its system prompt."
)

doc.add_heading("What It Does", level=2)
bullet("Answers questions about 4 mechanical keyboards ($109\u2013$199)")
bullet("Recommends products based on use case (gaming, typing, portability)")
bullet("Explains shipping options, return policy, and warranty")
bullet("Conversational only \u2014 users click \u201cAdd to Cart\u201d on the page themselves")

doc.add_heading("keyboard-demo/server.py", level=2)
doc.add_paragraph(
    "Same FastAPI relay pattern as Aria, but simpler \u2014 no tools. The system prompt contains "
    "the full product catalog (4 keyboards with specs, prices, and recommendations), shipping "
    "tiers, and conversation guidelines. Port 3001."
)

doc.add_heading("keyboard-demo/public/ \u2014 The Frontend", level=2)
doc.add_paragraph("Unlike Aria\u2019s single-file approach, this demo splits the frontend into separate files:")

add_table(
    ["File", "What It Does"],
    [
        ["index.html", "Page structure: nav, hero section, product grid, shipping banner, footer, voice widget mount point"],
        ["css/styles.css", "763 lines. Dark theme (#0a0a0f), purple accent (#6c5ce7), product cards with hover effects, mic button with pulse/spinner animations, responsive at 768px and 480px"],
        ["js/products.js", "Product data (window.PRODUCTS array, window.SHIPPING_INFO object) + DOM renderer that builds product cards on page load"],
        ["js/audio-recorder.js", "AudioRecorder class: getUserMedia at 16kHz, ScriptProcessorNode, Float32\u2192Int16\u2192base64. Echo cancellation + noise suppression enabled."],
        ["js/audio-streamer.js", "AudioStreamer class: decodes base64\u2192Int16\u2192Float32, schedules playback at 24kHz with look-ahead. clearQueue() for instant barge-in."],
        ["js/gemini-live-client.js", "WebSocket client connecting to ws://localhost:3001/ws. Sends {type: 'audio', data} messages. Exposes callbacks: onready, onaudio, onturncomplete, oninterrupted, etc."],
        ["js/voice-widget.js", "UI controller. Injects mic button into #voice-widget-container. State machine: idle \u2192 connecting \u2192 listening \u2192 agent_speaking. Wires all other JS classes together."],
    ],
)

doc.add_page_break()

# ===========================================================================
# 7. ARCHITECTURE DEEP DIVE
# ===========================================================================
doc.add_heading("7. Architecture Deep Dive", level=1)

doc.add_heading("Why a Server Relay (Not Direct Browser \u2192 Gemini)?", level=2)
bullet("The google-genai SDK is Python-only. Raw WebSocket from the browser has protocol quirks that cause silent failures and hangs.")
bullet("No browser-safe auth \u2014 you\u2019d expose your API key in page source.")
bullet("The SDK handles connection management and the Gemini protocol correctly.")
bullet("Server-side tool execution is more secure (API keys for ClickUp/Calendar stay on the server).")

doc.add_heading("Audio Pipeline (Step by Step)", level=2)
add_code(
    " 1. User speaks into microphone\n"
    " 2. Browser captures via getUserMedia() at 16kHz mono\n"
    " 3. Audio processor converts Float32 \u2192 Int16 PCM \u2192 base64\n"
    " 4. Browser sends {type: 'audio', data: '<base64>'} via WebSocket\n"
    " 5. Server decodes base64 \u2192 raw bytes\n"
    " 6. Server calls session.send_realtime_input(\n"
    "        audio=types.Blob(data=bytes, mime_type='audio/pcm;rate=16000')\n"
    "    )\n"
    " 7. Gemini processes audio natively (no STT step)\n"
    " 8. Gemini generates audio response natively (no TTS step)\n"
    " 9. Server receives audio from session.receive()\n"
    "10. Server base64-encodes \u2192 sends {type: 'audio', data: '<base64>'}\n"
    "11. Browser decodes base64 \u2192 Int16 \u2192 Float32, schedules playback at 24kHz\n"
    "12. User hears the response"
)

doc.add_heading("Barge-In (Interruption) Flow", level=2)
numbered("User starts speaking while agent audio is playing")
numbered("Gemini\u2019s built-in Voice Activity Detection (VAD) detects the user\u2019s voice")
numbered("Gemini stops generating and sends an \u2018interrupted\u2019 message")
numbered("Server relays {type: 'interrupted'} to browser")
numbered("Browser clears the audio playback queue (instant silence)")
numbered("UI state returns to \u2018listening\u2019 \u2014 ready for the user\u2019s new input")

doc.add_heading("WebSocket Protocol (Browser \u2194 Server)", level=2)
add_table(
    ["Direction", "Message", "Purpose"],
    [
        ["Browser \u2192 Server", '{"type": "audio", "data": "<base64>"}', "Mic audio chunk"],
        ["Server \u2192 Browser", '{"type": "audio", "data": "<base64>"}', "Agent audio response"],
        ["Server \u2192 Browser", '{"type": "status", "message": "Connected to Gemini"}', "Session ready"],
        ["Server \u2192 Browser", '{"type": "input_transcript", "text": "..."}', "What the user said"],
        ["Server \u2192 Browser", '{"type": "output_transcript", "text": "..."}', "What the agent said"],
        ["Server \u2192 Browser", '{"type": "turn_complete"}', "Agent finished speaking"],
        ["Server \u2192 Browser", '{"type": "interrupted"}', "User barged in"],
        ["Server \u2192 Browser", '{"type": "tool_call", "name": "..."}', "Tool being executed (Aria)"],
        ["Server \u2192 Browser", '{"type": "tool_result", "name": "...", "status": "..."}', "Tool finished (Aria)"],
        ["Server \u2192 Browser", '{"type": "error", "message": "..."}', "Error occurred"],
    ],
)

doc.add_heading("Critical Implementation Detail: send_realtime_input", level=2)
doc.add_paragraph(
    "When sending audio to Gemini, you MUST use the audio= parameter:"
)
add_code(
    "# CORRECT \u2014 works with Gemini 3.1 Flash Live\n"
    "await session.send_realtime_input(\n"
    "    audio=types.Blob(data=audio_bytes, mime_type='audio/pcm;rate=16000')\n"
    ")\n"
    "\n"
    "# WRONG \u2014 deprecated, Gemini 3.1 rejects it\n"
    "await session.send_realtime_input(\n"
    "    media=types.Blob(data=audio_bytes, mime_type='audio/pcm;rate=16000')\n"
    ")"
)

doc.add_page_break()

# ===========================================================================
# 8. SETUP GUIDE
# ===========================================================================
doc.add_heading("8. Setup Guide \u2014 From Zero to Running", level=1)

doc.add_heading("Prerequisites", level=2)
bullet("Python 3.10+")
bullet("A Google AI API key (free tier works) \u2014 https://aistudio.google.com/apikey")
bullet("A microphone")
bullet("Chrome or Edge browser")
bullet("For Aria: A ClickUp account + API key (https://app.clickup.com/settings/apps)")
bullet("For Aria\u2019s Calendar: Google Cloud OAuth credentials (Desktop app type)")

doc.add_heading("Step 1: Create .env", level=2)
doc.add_paragraph("In the Gemini Voice/ root directory:")
add_code(
    "GOOGLE_API_KEY=your-gemini-api-key\n"
    "CLICKUP_API_KEY=your-clickup-api-key    # only for Aria"
)

doc.add_heading("Step 2: Install Dependencies", level=2)
add_code(
    "# For the keyboard demo:\n"
    "cd keyboard-demo\n"
    "pip install -r requirements.txt\n"
    "\n"
    "# For Aria (includes Google Calendar OAuth libs):\n"
    "cd aria-assistant\n"
    "pip install -r requirements.txt"
)

doc.add_heading("Step 3: Google Calendar Setup (Aria Only)", level=2)
numbered("Go to Google Cloud Console \u2192 APIs & Services \u2192 Credentials")
numbered("Click + CREATE CREDENTIALS \u2192 OAuth Client ID \u2192 Desktop app")
numbered("Download the JSON file into the aria-assistant/ folder")
numbered("On first run, a browser window opens for Google login")
numbered("After login, token.json is saved automatically \u2014 no login needed again")

doc.add_heading("Step 4: Run a Demo", level=2)
add_code(
    "# Aria (port 8000):\n"
    "cd aria-assistant\n"
    "python server.py\n"
    "# Open http://localhost:8000\n"
    "\n"
    "# Keyboard demo (port 3001):\n"
    "cd keyboard-demo\n"
    "python server.py\n"
    "# Open http://localhost:3001"
)

doc.add_heading("Step 5: Talk to It", level=2)
numbered("Click the mic button")
numbered("Grant microphone permission when prompted")
numbered("Wait for \u2018Connected\u2019 / \u2018Listening...\u2019 status")
numbered("Start talking naturally")
numbered("Click the mic button again to stop")

doc.add_heading("Troubleshooting", level=2)
add_table(
    ["Problem", "Cause", "Fix"],
    [
        ["\u2018Connecting\u2019 spins forever", "Wrong model or API key", "Check server console. Verify GOOGLE_API_KEY."],
        ["\u2018Session ended\u2019 immediately", "Model doesn\u2019t support bidiGenerateContent", "Use gemini-3.1-flash-live-preview"],
        ["No audio playback", "Browser autoplay policy", "Must click mic button first (user gesture)"],
        ["Mic denied", "Browser blocked access", "Check browser settings, use localhost or HTTPS"],
        ["Echo / feedback", "Speakers picked up by mic", "Use headphones"],
        ["Port in use", "Another service running", "Change port in server.py or kill other process"],
        ["Calendar 403 error", "OAuth not set up", "Follow Step 3 above for Google Calendar setup"],
    ],
)

doc.add_page_break()

# ===========================================================================
# 9. HOW TO ADD YOUR OWN TOOLS
# ===========================================================================
doc.add_heading("9. How to Add Your Own Tools", level=1)

doc.add_paragraph(
    "The Aria demo shows the pattern for adding tools. To add a new tool (e.g., a weather API):"
)

doc.add_heading("Step 1: Create the Python function", level=2)
add_code(
    "# tools/weather.py\n"
    "import requests\n"
    "\n"
    "def get_weather(city: str) -> dict:\n"
    '    try:\n'
    '        resp = requests.get(f"https://api.weather.com/v1/{city}")\n'
    '        return {"status": "success", "temperature": resp.json()["temp"]}\n'
    '    except Exception as e:\n'
    '        return {"status": "error", "error_message": str(e)}'
)

doc.add_heading("Step 2: Register it in server.py", level=2)
add_code(
    '# Add to TOOLS dict:\n'
    'TOOLS = {\n'
    '    ...\n'
    '    "get_weather": weather.get_weather,\n'
    '}\n'
    '\n'
    '# Add to TOOL_DECLARATIONS list:\n'
    'TOOL_DECLARATIONS = [\n'
    '    ...\n'
    '    {\n'
    '        "name": "get_weather",\n'
    '        "description": "Gets current weather for a city.",\n'
    '        "parameters": {\n'
    '            "type": "object",\n'
    '            "properties": {\n'
    '                "city": {"type": "string", "description": "City name"}\n'
    '            },\n'
    '            "required": ["city"]\n'
    '        }\n'
    '    }\n'
    ']'
)

doc.add_heading("Step 3: Update the system prompt", level=2)
doc.add_paragraph(
    "Add a line to SYSTEM_PROMPT telling the model about the new capability: "
    "\u201cYou can check the weather for any city.\u201d Gemini uses the system prompt + tool "
    "declarations together to decide when to call tools."
)

doc.add_heading("That\u2019s It", level=2)
doc.add_paragraph(
    "No other changes needed. The execute_tool() function in server.py already handles "
    "dispatching any function in the TOOLS dict. Gemini will automatically start calling "
    "your new tool when the conversation warrants it."
)

doc.add_page_break()

# ===========================================================================
# 10. WHAT TO TELL CLAUDE CODE
# ===========================================================================
doc.add_heading("10. What to Tell Claude Code", level=1)

doc.add_paragraph(
    "If you want to build a voice agent from scratch using Claude Code, here\u2019s the prompt. "
    "Adjust the product/brand details and tools for your use case."
)

doc.add_heading("The Prompt", level=2)
add_code(
    'Build me a real-time voice agent using Google Gemini 3.1 Flash Live.\n'
    '\n'
    'Architecture:\n'
    '- Python FastAPI server as a WebSocket relay between browser and Gemini\n'
    '- Use the google-genai Python SDK (not raw WebSockets from the browser)\n'
    '- Model: gemini-3.1-flash-live-preview\n'
    '- Audio: PCM16, 16kHz mono input / 24kHz mono output\n'
    '- CRITICAL: Use send_realtime_input(audio=blob) not media=blob\n'
    '\n'
    'Server (FastAPI):\n'
    '- /ws WebSocket endpoint\n'
    '- Two async tasks: browser_to_gemini() and gemini_to_browser()\n'
    '- Run both with asyncio.gather()\n'
    '- Load GOOGLE_API_KEY from .env\n'
    '- Serve static files from public/\n'
    '- System prompt with [YOUR AGENT KNOWLEDGE]\n'
    '- [Optional: tool declarations + TOOLS dict + execute_tool()]\n'
    '\n'
    'Browser protocol:\n'
    '- Send: {type: "audio", data: "<base64 pcm16>"}\n'
    '- Receive: audio, status, turn_complete, interrupted, error\n'
    '\n'
    'Audio in browser:\n'
    '- Mic: getUserMedia at 16kHz, convert Float32->Int16->base64\n'
    '- Playback: decode base64->Int16->Float32, schedule at 24kHz\n'
    '- clearQueue() for instant barge-in (disconnect/reconnect gain node)\n'
    '- All AudioContext creation inside click handler (autoplay policy)'
)

doc.add_heading("Common Mistakes AI Agents Make (Without These Instructions)", level=2)
bullet("Connecting directly from browser to Gemini via raw WebSocket \u2014 unreliable, protocol issues")
bullet("Using the wrong model ID \u2014 must be one that supports bidiGenerateContent")
bullet("Using media= instead of audio= in send_realtime_input \u2014 deprecated, 3.1 rejects it")
bullet("Not handling barge-in \u2014 audio keeps playing when user tries to interrupt")
bullet("Creating AudioContext at page load \u2014 browsers block this, must be inside a click handler")
bullet("Putting systemInstruction in the setup message with audio-only responseModalities \u2014 known to hang on some models")

doc.add_page_break()

# ===========================================================================
# 11. DEPLOYMENT
# ===========================================================================
doc.add_heading("11. Deployment to Production", level=1)

doc.add_paragraph(
    "Both demos need a persistent server (not serverless) because the WebSocket stays open "
    "for the entire conversation. Vercel, Netlify, and Lambda won\u2019t work."
)

add_table(
    ["Platform", "Starting Price", "WebSocket Support", "Effort"],
    [
        ["Railway", "$5/mo (included credits)", "Yes", "Low \u2014 connect repo, done"],
        ["Render", "$7/mo", "Yes", "Low \u2014 connect repo, set start cmd"],
        ["Fly.io", "~$2-5/mo", "Yes", "Medium \u2014 needs Dockerfile"],
        ["DigitalOcean", "$4-6/mo (droplet)", "Yes", "Medium \u2014 manage yourself"],
    ],
)

doc.add_heading("Railway Deploy (Easiest)", level=2)
numbered("Push your project to GitHub")
numbered("Create a new Railway project from your repo")
numbered("Set environment variables: GOOGLE_API_KEY (and CLICKUP_API_KEY for Aria)")
numbered("Set start command: uvicorn server:app --host 0.0.0.0 --port $PORT")
numbered("Deploy \u2014 you get a public URL with HTTPS + WSS automatically")

doc.add_heading("Security Note", level=2)
doc.add_paragraph(
    "Add rate limiting or authentication to the /ws endpoint in production. Without it, "
    "anyone can open voice sessions and burn your API credits."
)

doc.add_page_break()

# ===========================================================================
# 12. KNOWN ISSUES
# ===========================================================================
doc.add_heading("12. Known Issues & Tips", level=1)

doc.add_heading("Known Gemini Live API Issues", level=2)
add_table(
    ["Issue", "Details", "Workaround"],
    [
        ["WebSocket 1011 errors", "Random timeouts during tool execution", "Retry; keep sessions under 10 min"],
        ["systemInstruction hang", "Setup hangs with audio-only + system prompt on some models", "Use google-genai SDK (handles it correctly)"],
        ["Model not found (1008)", "Model doesn\u2019t support bidiGenerateContent", "Use gemini-3.1-flash-live-preview"],
        ["No response to audio", "Some inputs get no response", "Ensure PCM16 at exactly 16kHz mono"],
        ["Config locked after connect", "Can\u2019t change model/tools/prompt mid-session", "Disconnect and reconnect"],
    ],
)

doc.add_heading("Tips", level=2)
bullet("Use headphones during testing to avoid echo/feedback.")
bullet("Browser DevTools Console shows all WebSocket messages \u2014 check there first when debugging.")
bullet("Server terminal shows Gemini SDK errors \u2014 watch both browser and server logs.")
bullet("Sessions expire after 15 minutes. Click mic again to reconnect.")
bullet("The system prompt is in server.py. Edit it to change personality or knowledge.")
bullet("To change the voice, modify the speechConfig in LiveConnectConfig. Voices: Kore, Puck, Charon, Fenrir, Aoede.")
bullet("For lowest latency, deploy near Google\u2019s API endpoints (us-central1).")
bullet("The Aria demo generates the system prompt at import time (captures current date/time). Restart the server to update the time context.")
bullet("All tool functions are synchronous (using requests). For slow APIs, consider running them in a thread executor to avoid blocking.")

# ===========================================================================
# SAVE
# ===========================================================================
output_path = os.path.join(
    os.path.dirname(__file__), "..",
    "Gemini Voice Widget - Master Resource Guide.docx",
)
doc.save(output_path)
print(f"Saved to: {os.path.abspath(output_path)}")
